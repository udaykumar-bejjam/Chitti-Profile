#!/usr/bin/env python3
"""Generate profile PDF from the Word template to preserve exact formatting."""

from __future__ import annotations

import shutil
import subprocess
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "chinni_resume_fresher.docx"
OUTPUT_DOCX = ROOT / "jessika_resume.docx"
OUTPUT_PDF = ROOT / "Jessika_Seedarla_Profile.pdf"

CORE_SKILLS = [
    "PCR",
    "ELISA",
    "Chromatography",
    "Electrophoresis",
    "Gram staining",
    "Sterilization techniques",
    "DNA isolation",
    "Blotting techniques",
    "Replication of DNA",
    "Transcription",
    "Translation",
    "Data documentation",
]


def set_run_text(run, text: str) -> None:
    run.text = text


def set_bullet(paragraph, label: str, description: str) -> None:
    runs = paragraph.runs
    if len(runs) >= 4:
        set_run_text(runs[2], label)
        set_run_text(runs[3], description)
    elif len(runs) == 1:
        set_run_text(runs[0], f"•\t{label}{description}")


def set_skill_bullet(paragraph, skill: str) -> None:
    runs = paragraph.runs
    if len(runs) >= 4:
        set_run_text(runs[2], skill)
        set_run_text(runs[3], "")
    elif len(runs) == 1:
        set_run_text(runs[0], f"•\t{skill}")


def set_body_parts(paragraph, parts: list[tuple[str, bool]]) -> None:
    runs = paragraph.runs
    for idx, (text, _) in enumerate(parts):
        if idx < len(runs):
            set_run_text(runs[idx], text)
    if len(runs) > len(parts):
        for run in runs[len(parts) :]:
            set_run_text(run, "")


def delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def find_paragraph(doc: Document, text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text}")


def set_core_competencies(doc: Document) -> None:
    heading_idx = find_paragraph(doc, "Core Technical Competencies")
    strengths_idx = find_paragraph(doc, "Personal Strengths")

    bullet_start = heading_idx + 1
    existing_bullets = strengths_idx - bullet_start

    for offset in range(existing_bullets):
        set_skill_bullet(doc.paragraphs[bullet_start + offset], CORE_SKILLS[offset])

    anchor = doc.paragraphs[bullet_start]
    for skill in CORE_SKILLS[existing_bullets:]:
        new_p = deepcopy(anchor._element)
        doc.paragraphs[strengths_idx - 1]._element.addnext(new_p)
        set_skill_bullet(Paragraph(new_p, anchor._parent), skill)
        strengths_idx += 1


def compact_for_two_pages(doc: Document) -> None:
    """Keep the resume within two pages after the expanded core skills section."""
    paragraphs = doc.paragraphs

    for text in (
        "Career Objective",
        "Education",
        "Academic Projects & Laboratory Training",
        "Certifications & Training",
        "Laboratory Skills",
        "Core Technical Competencies",
        "Personal Strengths",
        "Languages",
        "Additional Information",
    ):
        paragraphs[find_paragraph(doc, text)].paragraph_format.space_before = Pt(7)

    paragraphs[find_paragraph(doc, "Career Summary")].paragraph_format.space_after = Pt(3)
    paragraphs[2].paragraph_format.space_after = Pt(7)

    for text in ("Career Summary", "Career Objective"):
        body = paragraphs[find_paragraph(doc, text) + 1]
        body.paragraph_format.space_after = Pt(3)
        body.paragraph_format.line_spacing = 1.08

    core_start = find_paragraph(doc, "Core Technical Competencies") + 1
    strengths_idx = find_paragraph(doc, "Personal Strengths")
    for idx in range(core_start, strengths_idx):
        bullet = paragraphs[idx]
        bullet.paragraph_format.space_after = Pt(0)
        bullet.paragraph_format.line_spacing = 1.0

    for text in ("Key Attributes:", "Areas of Interest:"):
        for paragraph in paragraphs:
            if paragraph.text.strip().startswith(text):
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05


def build_document() -> Document:
    doc = Document(TEMPLATE)
    p = doc.paragraphs

    set_run_text(p[0].runs[0], "Jessika ")
    set_run_text(p[0].runs[1], "Seedarla")
    set_run_text(
        p[1].runs[0],
        "M.Sc. Biotechnology (2026)  ·  Gold Medalist  ·  Fresher  ·  Entry-Level Biopharma Candidate",
    )

    set_body_parts(
        p[4],
        [
            ("Recent ", False),
            ("M.Sc. Biotechnology graduate (2026) and Gold Medalist", True),
            (" seeking an opportunity in biopharma and life sciences. Brings a strong academic record as a ", False),
            ("Rank 1 B.Sc. Biotechnology graduate", True),
            (
                " and practical laboratory exposure through dissertation research, academic projects, and analytical laboratory training. Motivated to learn industrial processes, work with SOPs, and grow in QC microbiology and bioprocess support.",
                False,
            ),
        ],
    )

    set_run_text(p[8].runs[0], "M.Sc. Biotechnology – Krishna University, Machilipatnam (Gold Medal)")

    new_p = deepcopy(p[21]._element)
    p[20]._element.addnext(new_p)
    gold_para = Paragraph(new_p, p[20]._parent)
    set_bullet(
        gold_para,
        "Gold Medal in M.Sc. Biotechnology – ",
        "Awarded for outstanding academic performance at Krishna University, Machilipatnam",
    )

    p = doc.paragraphs
    set_bullet(
        p[22],
        "Ranked 1st in B.Sc. Biotechnology – ",
        "Consistently topped all semesters at Harshini Degree College",
    )
    set_bullet(
        p[23],
        "Top Scorer in Intermediate (BiPC) – ",
        "Distinction grades with special excellence in Biology",
    )
    set_bullet(
        p[24],
        "Top Scorer in Class 10 (SSC) – ",
        "Distinction with outstanding performance in Science",
    )
    set_bullet(
        p[25],
        "Merit Certificates & Scholarships – ",
        "Received multiple honors for academic excellence",
    )
    set_bullet(
        p[26],
        "Perfect Academic Record – ",
        "Maintained consistency throughout school and college",
    )

    p = doc.paragraphs
    new_p = deepcopy(p[34]._element)
    p[33]._element.addnext(new_p)
    analytical_para = Paragraph(new_p, p[33]._parent)
    set_bullet(
        analytical_para,
        "Analytical Basics – ",
        "Introductory HPLC exposure and strong academic grounding in scientific analysis",
    )

    set_core_competencies(doc)
    compact_for_two_pages(doc)
    return doc


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    outdir = pdf_path.parent
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            str(docx_path),
            "--outdir",
            str(outdir),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    generated = outdir / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        generated.replace(pdf_path)


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT_DOCX)
    export_pdf(OUTPUT_DOCX, OUTPUT_PDF)

    for name in ("profile.pdf", "jessika_resume.pdf", "jessika_resume_print.pdf"):
        shutil.copy2(OUTPUT_PDF, ROOT / name)

    chitti = ROOT / "chitti-profile"
    if chitti.is_dir():
        shutil.copy2(OUTPUT_DOCX, chitti / "jessika_resume.docx")
        for name in (
            "Jessika_Seedarla_Profile.pdf",
            "profile.pdf",
            "jessika_resume.pdf",
            "jessika_resume_print.pdf",
        ):
            shutil.copy2(OUTPUT_PDF, chitti / name)

    print(f"Wrote {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
